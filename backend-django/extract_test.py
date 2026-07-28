import os
from docx import Document
p = r'C:\Users\Dell\Desktop\minor project\originality_guard\backend-django\media\submissions\about_final.docx'
print('EXISTS', os.path.exists(p))
if os.path.exists(p):
    print('SIZE', os.path.getsize(p))
try:
    doc = Document(p)
    print('PARA COUNT', len(doc.paragraphs))
    for i, para in enumerate(doc.paragraphs[:20]):
        print(i, repr(para.text))
except Exception as e:
    import traceback
    traceback.print_exc()
